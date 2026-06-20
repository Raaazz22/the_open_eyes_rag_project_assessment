from docx import Document as DocxDocument
from langchain_core.documents import Document


def word_extractor(file_path,file_name):

    extracted_text = DocxDocument(file_path)
    docs = []
    for p in extracted_text.paragraphs:
        doc = Document(
            page_content=p.text,
            metadata={"file_name": file_name}
        )
        docs.append(doc)
    return docs

file_path = "sample_docx.docx"
# results = word_extractor(file_path,"xyz")
# print(results)
